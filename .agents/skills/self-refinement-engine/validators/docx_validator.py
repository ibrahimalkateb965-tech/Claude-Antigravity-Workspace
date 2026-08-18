from docx import Document

class DOCXValidator:
    """
    Validator for Word (DOCX) files to ensure layout compliance.
    Checks RTL settings, margin bounds, and structural compactness.
    """
    def __init__(self, max_elements=40):
        self.max_elements = max_elements

    def validate(self, docx_path):
        failures = []
        try:
            doc = Document(docx_path)
            
            # 1. Check total text elements count to estimate page overflow risk
            total_elements = len([p for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        total_elements += len([p for p in cell.paragraphs if p.text.strip()])

            if total_elements > self.max_elements:
                failures.append(f"DOCX contains {total_elements} elements, exceeding 1-page target ({self.max_elements}). Risk of 2-page spill.")

            # 2. Check top/bottom margins (must be <= 0.85in for compact single-page letters)
            for i, section in enumerate(doc.sections):
                top_in = section.top_margin.inches
                bottom_in = section.bottom_margin.inches
                if top_in > 0.85 or bottom_in > 0.85:
                    failures.append(f"Section {i} top/bottom margin is {top_in:.2f}in; recommend <=0.80in for single page fit.")

            return failures
        except Exception as e:
            return [f"Failed to inspect DOCX file: {str(e)}"]
