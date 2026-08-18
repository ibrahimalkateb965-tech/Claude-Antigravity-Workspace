import os
import base64
from io import BytesIO
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
import contextlib

def apply_rtl_font(run, size_pt, bold=False, italic=False):
    """Applies strict Arabic font styling to a run."""
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    
    # Override for Arabic (Complex Scripts)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:cs'), 'Arial')
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')
    
    # Set Complex Script size (w:szCs)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)
    
    # RTL run property
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    
    if bold:
        bCs = OxmlElement('w:bCs')
        bCs.set(qn('w:val'), '1')
        rPr.append(bCs)

def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Forces paragraph to be RTL and sets alignment safely based on MEMORY_STORE.md."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    
    # LESSON-DOCX-011: Avoid WD_ALIGN_PARAGRAPH.RIGHT due to MS Word rendering bug on some versions.
    if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
    elif align in [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]:
        paragraph.alignment = align
    # If align is RIGHT, we do nothing to <w:jc>, because bidi naturally aligns to the right.

def apply_list_rtl(paragraph, indent_inches=0.25):
    """Forces right-aligned list bullet/number indentation."""
    set_paragraph_rtl(paragraph, None) # Naturally right-aligned by bidi
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.get_or_add_ind()
    # Remove left indentation
    ind.set(qn('w:left'), "0")
    # Add right indentation for RTL
    ind.set(qn('w:right'), str(int(indent_inches * 1440)))
    ind.set(qn('w:hanging'), str(int(0.25 * 1440)))

def enable_auto_update_toc(doc):
    """Injects w:updateFields into settings.xml to auto-update TOC on open."""
    settings = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)

def process_inline_nodes(nodes, paragraph, default_size=14):
    """Processes mistune inline nodes (text, strong, emphasis, linebreak)."""
    if not nodes:
        return
    for node in nodes:
        node_type = node.get('type')
        if node_type == 'text':
            run = paragraph.add_run(node.get('raw', ''))
            apply_rtl_font(run, default_size)
        elif node_type == 'strong':
            text = ''.join(c.get('raw', '') for c in node.get('children', []) if c.get('type') == 'text')
            run = paragraph.add_run(text)
            apply_rtl_font(run, default_size, bold=True)
        elif node_type == 'emphasis':
            text = ''.join(c.get('raw', '') for c in node.get('children', []) if c.get('type') == 'text')
            run = paragraph.add_run(text)
            apply_rtl_font(run, default_size, italic=True)
        elif node_type == 'linebreak' or node_type == 'softbreak':
            run = paragraph.add_run()
            run.add_break()
        elif node_type == 'image':
            url = node.get('attrs', {}).get('url', '')
            try:
                if url.startswith('data:image'):
                    # data:image/png;base64,iVBORw0KGgo...
                    header, encoded = url.split(',', 1)
                    image_data = base64.b64decode(encoded)
                    image_stream = BytesIO(image_data)
                    run = paragraph.add_run()
                    run.add_picture(image_stream, width=Inches(6.0))
                elif os.path.isabs(url) and os.path.exists(url):
                    run = paragraph.add_run()
                    run.add_picture(url, width=Inches(6.0))
            except Exception as e:
                pass # Graceful fail for images
        else:
            # Fallback for nested elements or unknown
            if 'children' in node:
                process_inline_nodes(node['children'], paragraph, default_size)
            elif 'raw' in node:
                run = paragraph.add_run(node['raw'])
                apply_rtl_font(run, default_size)

def walk_ast_and_build(ast, placeholder_p, doc):
    """Walks the mistune AST and constructs Word elements before placeholder_p."""
    for node in ast:
        try:
            node_type = node.get('type')
            
            if node_type == 'heading':
                level = node.get('attrs', {}).get('level', 1)
                # Map to Heading 1, Heading 2...
                style_name = f'Heading {level}' if level <= 9 else 'Heading 1'
                try:
                    p = placeholder_p.insert_paragraph_before(style=style_name)
                except KeyError:
                    p = placeholder_p.insert_paragraph_before()
                set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.RIGHT)
                # Clear default runs from style if any, though insert_paragraph_before creates empty
                process_inline_nodes(node.get('children', []), p, default_size=18)
                
            elif node_type == 'paragraph':
                p = placeholder_p.insert_paragraph_before()
                set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
                
                # Check for preamble (ends with colon)
                is_preamble = False
                children = node.get('children', [])
                if children:
                    last_child = children[-1]
                    if last_child.get('type') == 'text' and last_child.get('raw', '').strip().endswith(':'):
                        is_preamble = True
                
                # LESSON-DOCX-009 & LESSON-DOCX-010: Apply first line indent EXCEPT for preambles
                if not is_preamble:
                    pPr = p._p.get_or_add_pPr()
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:firstLine'), '360')
                    pPr.append(ind)

                process_inline_nodes(children, p, default_size=14)
                
            elif node_type == 'list':
                is_ordered = node.get('attrs', {}).get('ordered', False)
                for item in node.get('children', []):
                    if item.get('type') == 'list_item':
                        # Add paragraph for list item
                        try:
                            style = 'List Number' if is_ordered else 'List Bullet'
                            p = placeholder_p.insert_paragraph_before(style=style)
                        except KeyError:
                            p = placeholder_p.insert_paragraph_before()
                        apply_list_rtl(p, indent_inches=0.25)
                        # Mistune list_item children are usually block_text
                        for child in item.get('children', []):
                            if child.get('type') == 'block_text':
                                process_inline_nodes(child.get('children', []), p, default_size=14)
                            elif child.get('type') == 'list':
                                # Basic nested list support (recursive could be complex with python-docx styles, keep it flat for now)
                                pass
                                
            elif node_type == 'table':
                # Create table at the end, then move it before placeholder
                # Extract headers and body to calculate rows/cols
                head = next((c for c in node.get('children', []) if c.get('type') == 'table_head'), None)
                body = next((c for c in node.get('children', []) if c.get('type') == 'table_body'), None)
                
                rows_data = []
                alignments = []
                
                if head:
                    cells = head.get('children', [])
                    rows_data.append(cells)
                    # Mistune 3.x attrs might be None if no alignment specified
                    for cell in cells:
                        attrs = cell.get('attrs')
                        alignments.append(attrs.get('align') if attrs else None)
                
                if body:
                    for row_node in body.get('children', []):
                        rows_data.append(row_node.get('children', []))
                        
                if not rows_data:
                    continue
                    
                cols = len(rows_data[0])
                rows = len(rows_data)
                
                # Explicit Table Grid to ensure borders are visible
                try:
                    tbl = doc.add_table(rows=rows, cols=cols, style='Table Grid')
                except KeyError:
                    tbl = doc.add_table(rows=rows, cols=cols)
                
                # Make Table RTL
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tblPr = tbl._tbl.tblPr
                bidiVisual = OxmlElement('w:bidiVisual')
                tblPr.append(bidiVisual)
                
                # Populate cells
                for r_idx, row_node in enumerate(rows_data):
                    for c_idx, cell_node in enumerate(row_node):
                        if c_idx >= cols:
                            break
                        cell = tbl.cell(r_idx, c_idx)
                        p = cell.paragraphs[0]
                        
                        # Reverse alignment for RTL
                        align = alignments[c_idx] if c_idx < len(alignments) else None
                        wd_align = WD_ALIGN_PARAGRAPH.RIGHT # RTL Default Fallback
                        if align == 'left': wd_align = WD_ALIGN_PARAGRAPH.RIGHT # RTL inverse
                        elif align == 'right': wd_align = WD_ALIGN_PARAGRAPH.LEFT # RTL inverse
                        elif align == 'center': wd_align = WD_ALIGN_PARAGRAPH.CENTER
                        
                        set_paragraph_rtl(p, wd_align)
                        process_inline_nodes(cell_node.get('children', []), p, default_size=14)
                
                # Move table XML before placeholder paragraph
                placeholder_p._p.addprevious(tbl._tbl)
        except Exception as e:
            import sys
            print(f"Error processing node {node.get('type', 'unknown')}: {str(e)}", file=sys.stderr)
            continue

@contextlib.contextmanager
def safe_docx_edit(template_path, output_path):
    """Context manager for safely opening and saving a DOCX file without locking."""
    doc = docx.Document(template_path)
    try:
        yield doc
    finally:
        doc.save(output_path)

def build_docx_from_ast(ast, template_path, output_path):
    with safe_docx_edit(template_path, output_path) as doc:
        enable_auto_update_toc(doc)
        
        # Find placeholder
        placeholder_p = None
        for p in doc.paragraphs:
            if '{{CONTENT_HERE}}' in p.text:
                placeholder_p = p
                break
                
        if not placeholder_p:
            # Fallback if placeholder not found: just append to the end
            placeholder_p = doc.add_paragraph()
            
        walk_ast_and_build(ast, placeholder_p, doc)
        
        # Erase the placeholder text carefully without deleting the paragraph (preserves section properties)
        for run in placeholder_p.runs:
            run.text = run.text.replace('{{CONTENT_HERE}}', '')
